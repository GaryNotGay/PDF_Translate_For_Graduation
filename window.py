# @Author  :  lijishi
# @Contact :  lijishi@emails.bjut.edu.cn
# @Software:  Pycharm
# @EditTime:  Jan 31,2020
# @describe:  GUI of PDF Translate
# @LICENSE :  GNU GENERAL PUBLIC LICENSE Version 3

# This is only for graduation design
# Lack of function and update slowly

# fitz pillow python-docx pdfminer pdfminer3k pdfminer.six tencent PyMuPDF request

import os
import time
import fitz
import datetime
import threading
import tkinter as tk
import tkinter.filedialog
from docx import Document
from docx.oxml.ns import qn
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter import scrolledtext
from PIL import Image, ImageTk
from pdf2txt import parse
from txtArrange import Arrange
from YoudaoTrans import connect

class MyThread(threading.Thread):
    def __init__(self,func,args=()):
        super(MyThread,self).__init__()
        self.func = func
        self.args = args
    def run(self):
        self.result = self.func(*self.args)
    def get_result(self):
        try:
            return self.result
        except Exception:
            return None

def SelectPath_In():
    global page_num
    path_ = tkinter.filedialog.askopenfilename()
    path_ = path_.replace("/", "\\\\")
    path_in.set(path_)
    string = path_
    i = string.rfind("\\\\")
    string = string[:i+2]
    path_out.set(string)
    page_num = 0
    PDF_View(path_, page_num)

def SelectPath_Out():
    path_ = tkinter.filedialog.askdirectory()
    path_ = path_.replace("/", "\\\\")
    path_out.set(path_)

def Last_Page():
    global page_num
    if page_num == 0:
        tk.messagebox.showerror("Error", "已至首页")
        return
    page_num -= 1
    PDF_View(path_in.get(), page_num)

def Next_Page():
    global page_num
    global total_num
    if page_num == total_num-1:
        tk.messagebox.showerror("Error", "已至尾页")
        return
    page_num += 1
    PDF_View(path_in.get(), page_num)

def Go_Page():
    global page_num
    global total_num
    page_num = go_num.get() - 1
    if page_num > total_num-1:
        tk.messagebox.showerror("Error", "页码错误")
        return
    if page_num < 0:
        tk.messagebox.showerror("Error", "页码错误")
        return
    PDF_View(path_in.get(), page_num)

def PDF_View(path, page_num):
    global total_num
    doc = fitz.open(path)
    page = doc[page_num]
    total_num = doc.pageCount
    out_num = str(page_num+1) + r'/' + str(total_num)
    cur_total.set(out_num)
    pageview = page.getPixmap(alpha=False)
    pageview.writeImage("pdf_view_temp.jpg")
    image = Image.open("pdf_view_temp.jpg")
    img_temp = image.resize((440, 760), Image.ANTIALIAS)
    img = ImageTk.PhotoImage(img_temp)
    view.configure(image=img)
    os.remove("pdf_view_temp.jpg")
    main_window.update_idletasks()
    main_window.mainloop()

def Tips():
    tk.messagebox.showinfo("Tips", "请注意PDF文件大小，过大则会影响翻译速度\n翻译速度受网络状况等多重影响，请耐心等待\n本软件仅用于毕业设计，请勿挪用于非法用途\n若有道智云更改政策，可能导致软件无法使用")

def About():
    # window centered
    about_window = Toplevel()
    screen_width = about_window.winfo_screenwidth()
    screen_heigh = about_window.winfo_screenheight()
    about_window_width = 370
    about_window_heigh = 210
    x = (screen_width - about_window_width) / 2
    y = (screen_heigh - about_window_heigh) / 2
    about_window.geometry("%dx%d+%d+%d" % (about_window_width, about_window_heigh, x, y))

    # window layout
    global fan_gif
    about_window.title('About')
    about_window.iconbitmap(".\\picture\\fan.ico")
    fan_gif = tk.PhotoImage(file=".\\picture\\fan.gif")
    software_frame = ttk.LabelFrame(about_window, text='Software Info')
    software_frame.grid(row=0, column=0, rowspan=5, columnspan=4, padx=50, pady=5)
    ttk.Label(software_frame, image=fan_gif, compound='left').grid(row=0, rowspan=4, column=0)
    ttk.Label(software_frame, text="PDF Translate For Graduation").grid(row=0, column=1, sticky = W)
    ttk.Label(software_frame, text="@Author    :   lijishi").grid(row=1, column=1, sticky = W)
    ttk.Label(software_frame, text="@College   :   BJUT CS").grid(row=2, column=1, sticky=W)
    ttk.Label(software_frame, text="@EditTime  :   Jan 31,2020").grid(row=3, column=1, sticky=W)

    copyright_frame = ttk.LabelFrame(about_window, text='LICENSE Info')
    copyright_frame.grid(row=5, column=0, rowspan=3, columnspan=4, padx=50, pady=5)
    ttk.Label(copyright_frame, text = "Github @ PDF_Translate_For_Graduation").grid(row=5, column=0)
    ttk.Label(copyright_frame, text="GNU GENERAL PUBLIC LICENSE Version 3").grid(row=6, column=0)

def youdao(language):
    try:
        total_lines = 0
        for index, line in enumerate(open('.\\pdf2txt_temp.txt', 'r', encoding='utf-8')):
            total_lines += 1

        ready_txt = open('.\\pdf2txt_temp.txt', encoding='utf-8')
        for index in range(total_lines):
            line = ready_txt.readline()
            if line == '\n':
                continue
            elif line == ' \n':
                continue

            return_text = connect(line, language, int(double_language.get()))
            if return_text[0:11] == 'YouDaoError':
                tk.messagebox.showerror("YouDaoError", return_text[11:])
                now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
                scr.insert(END, now_time)
                scr.insert(END, ' 翻译失败\n')
                scr.insert(END, '\n')
                scr.insert(END, return_text[11:])
                ready_txt.close()
                os.remove('.\\pdf2txt_temp.txt')
                os.remove('.\\out_temp.txt')
                return
            with open('.\\out_temp.txt', 'a', encoding='utf-8') as out_txt:
                out_txt.write(return_text + '\n')
            out_txt.close()
        ready_txt.close()
        success = ' 翻译成功\n'
        later = '三秒后显示翻译文本\n'
        now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
        scr.insert(END, now_time)
        scr.insert(END, success)
        scr.insert(END, '\n')
        scr.insert(END, later)
        time.sleep(3)
        os.remove('.\\pdf2txt_temp.txt')
        OutFile()
    except BaseException as err:
        tk.messagebox.showerror("UnknownError", err)
        now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
        scr.insert(END, now_time)
        scr.insert(END, ' 翻译失败\n')
        scr.insert(END, '\n')
        scr.insert(END, err)
        out_txt.close()
        ready_txt.close()
        os.remove('.\\pdf2txt_temp.txt')
        os.remove('.\\out_temp.txt')
        return

def OutFile():
    global out_file
    global api_choose
    global language_change

    now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    scr.delete(1.0, END)
    scr.insert("insert", now_time)
    scr.insert("insert", " 文档生成成功")
    scr.insert("insert", "\n")
    scr.insert("insert", "\n")
    log_name = '.\\log.txt'
    with open(log_name, 'a', encoding='utf-8') as log:
        now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        log.write("Time ：")
        log.write(now_time)
        log.write("\n")
        log.write("File Name : ")
        log.write(str(path_in.get()))
        log.write("\n")
        log.write("Translate Mode : ")
        log.write(OutMode(api_choose, language_change))
        if out_file == 1:
            log.write(" TXT")
        elif out_file == 2:
            log.write(" DOC")
        if int(double_language.get()):
            log.write(" DoubleLanguage")
        log.write("\n")
        log.write("\n")
    log.close()
    total_lines = 0
    for index, line in enumerate(open('.\\out_temp.txt', 'r', encoding='utf-8')):
        total_lines += 1

    string = str(path_in.get())
    string = string.replace("\\", "\\\\")
    string = string.replace("/", "\\\\")
    i = string.rfind("\\\\")
    out_name_temp = string[i + 2:len(string) - 4]
    if out_file == 1:
        out_name = str(path_out.get()) + out_name_temp + "_Translate_" + str(int(time.time())) + ".txt"
        with open(out_name, 'a', encoding='utf-8') as out:
            now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
            out.write("-------------------------------------------------------\n")
            out.write("Created By PDF Translate At ")
            out.write(now_time)
            out.write("\n")
            out.write("Translate Mode : ")
            out.write(OutMode(api_choose, language_change))
            out.write("\n")
            out.write("-------------------------------------------------------\n")
            out.write("\n")
        out.close()
    elif out_file == 2:
        out_name = str(path_out.get()) + out_name_temp + "_Translate_" + str(int(time.time())) + ".doc"
        now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("------------------------------------------------------------\n")
        run = paragraph.add_run("Created By PDF Translate At ")
        run = paragraph.add_run(now_time)
        run = paragraph.add_run("\n")
        run = paragraph.add_run("Translate Mode : ")
        run = paragraph.add_run(OutMode(api_choose, language_change))
        run = paragraph.add_run("\n")
        run = paragraph.add_run("------------------------------------------------------------\n")
        run = paragraph.add_run("\n")
        font = run.font

    out_text = ''
    ready_txt = open('.\\out_temp.txt', encoding='utf-8')
    for index in range(total_lines):
        line = ready_txt.readline()
        scr.insert("insert", line)
        scr.insert("insert", "\n")
        if out_file == 1:
            with open(out_name, 'a', encoding='utf-8') as out:
                out.write(line)
                #out.write("\n")
            out.close()
        elif out_file == 2:
            run = paragraph.add_run(line)
            #run = paragraph.add_run("\n")

    if out_file == 2:
        doc.save(out_name)
    ready_txt.close()
    os.remove('.\\out_temp.txt')

def OutMode(api, language):
    global page_start
    global page_end
    out_text = ''

    if api == 1:
        out_text += '有道翻译 '
    elif api == 2:
        out_text += '百度翻译 '
    elif api == 2:
        out_text += '腾讯翻译 '

    if language == 1:
        out_text += 'ZH->EN '
    elif language == 2:
        out_text += 'EN->ZH '

    if page_start == 0 and page_end == 0:
        out_text += 'All Page'
    elif page_start != 0 and page_end == 0:
        out_text += 'Page '
        out_text += str(page_start)
    elif page_start != 0 and page_end != 0:
        out_text += 'Page '
        out_text += str(page_start)
        out_text += 'to '
        out_text += str(page_end)

    return out_text

def Translate():
    global out_file
    global api_choose
    global language_change
    global page_start
    global page_end

    if int(page.get()) == 1:
        page_start = 0
        page_end = 0
    elif int(page.get()) == 2:
        page_start = int(num_only.get())
        page_end = 0
    elif int(page.get()) == 3:
        page_start = int(num_start.get())
        page_end = int(num_end.get())

    if os.path.exists('.\\pdf2txt_temp.txt'):
        os.remove('.\\pdf2txt_temp.txt')
    if os.path.exists('.\\out_temp.txt'):
        os.remove('.\\out_temp.txt')
    scr.delete(1.0, END)
    now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    scr.insert("insert", now_time)
    scr.insert("insert", " 开始提取文字\n")
    parse(path_in.get(), '.\\pdf2txt_temp.txt', page_start, page_end)
    now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')
    scr.insert("insert", now_time)
    scr.insert("insert", " 文字提取成功\n")
    scr.insert("insert", now_time)
    scr.insert("insert", " 开始翻译...\n")

    #Because Unsolvable Bugs, Temporarily Disable txtArrange
    #Arrange('.\\pdf2txt_temp.txt', '.\\txtArrange_temp.txt')
    #os.remove('.\\pdf2txt_temp.txt')

    if language.get() == 'ZH->EN':
        language_change = 1
    elif language.get() == 'EN->ZH':
        language_change = 2
    else:
        tk.messagebox.showerror("Error","参数错误")

    if api.get() == '有道翻译':
        api_choose = 1
    elif api.get() == '百度翻译':
        api_choose = 2
    elif api.get() == '腾讯翻译':
        api_choose = 3
    else:
        tk.messagebox.showerror("Error", "参数错误")

    if file.get() == 'TXT':
        out_file = 1
    elif file.get() == 'DOC':
        out_file = 2
    else:
        tk.messagebox.showerror("Error", "参数错误")

    if api_choose == 1:
        thread_YouDao = MyThread(youdao, args=(language_change,))
        thread_YouDao.start()
    elif api_choose == 2:
        thread_Baidu = MyThread(baidu, args=(language_change,))
        thread_Baidu.start()
    elif api_choose == 3:
        thread_Tencent = MyThread(tencent, args=(language_change,))
        thread_Tencent.start()

# window centered
main_window=tk.Tk()
screen_width = main_window.winfo_screenwidth()
screen_heigh = main_window.winfo_screenheight()
main_window_width = 920
main_window_heigh = 850
x = (screen_width-main_window_width) / 2
y = (screen_heigh-main_window_heigh) / 2
main_window.geometry("%dx%d+%d+%d" %(main_window_width,main_window_heigh,x,y))

# window layout
main_window.title("PDF Translate For Graduation")
main_window.iconbitmap(".\\picture\\fan256.ico")
path_frame = ttk.LabelFrame(main_window, text = '文件路径选择')
path_frame.grid(row = 0, column = 0, rowspan = 2, columnspan = 4, padx=10, pady=5)
path_in = tk.StringVar()
path_out = tk.StringVar()
path_in.set("请选择源文件位置，也可在框中键入")
path_out.set("请选择输出位置，默认为源文件路径")
ttk.Label(path_frame, text = "源文件位置").grid(row = 0, column = 0, padx=10)
ttk.Entry(path_frame, width = 30, textvariable = path_in).grid(row = 0, column = 1, padx=5)
ttk.Button(path_frame, text = "选择", command = SelectPath_In).grid(row = 0, column = 2, padx=10)
ttk.Label(path_frame, text = "输出位置").grid(row = 1, column = 0, padx=10)
ttk.Entry(path_frame, width = 30, textvariable = path_out).grid(row = 1, column = 1, padx=5)
ttk.Button(path_frame, text = "选择", command = SelectPath_Out).grid(row = 1, column = 2, padx = 10, pady = 10)
page_frame = ttk.LabelFrame(main_window, text = '翻译页面选择')
page_frame.grid(row = 2, column = 0, rowspan = 3, columnspan = 4, padx=10, pady=5, sticky = W)
page = IntVar()
Radiobutton(page_frame, text = "全部翻译", variable = page, value = 1).grid(row = 2, column = 0, padx = 10, pady = 5)
Radiobutton(page_frame, text = "单页翻译", variable = page, value = 2).grid(row = 3, column = 0, padx = 10)
Radiobutton(page_frame, text = "多页翻译", variable = page, value = 3).grid(row = 4, column = 0, padx = 10)
ttk.Label(page_frame, text = "翻译页数：").grid(row = 3, column = 1)
num_only = tk.IntVar()
start_entry = ttk.Entry(page_frame, width = 5, textvariable = num_only).grid(row = 3, column = 2)
ttk.Label(page_frame, text = "翻译页数：").grid(row = 4, column = 1)
num_start = tk.IntVar()
num_end = tk.IntVar()
start_entry = ttk.Entry(page_frame, width = 5, textvariable = num_start).grid(row = 4, column = 2, padx = 10, pady = 10)
ttk.Label(page_frame, text = "—").grid(row = 4, column = 3)
end_entry = ttk.Entry(page_frame, width = 5, textvariable = num_end).grid(row = 4, column = 4, padx = 10, pady = 10)
outmode_frame = ttk.LabelFrame(main_window, text = '翻译类型选择')
outmode_frame.grid(row = 2, column = 3, rowspan = 3, padx=10, pady=5, sticky = E)
language = tk.StringVar()
languageChosen = ttk.Combobox(outmode_frame, width=8, textvariable=language)
languageChosen.grid(row = 2, column = 5, padx = 5, pady = 6, sticky = W)
languageChosen['values'] = ('语种选择', 'ZH->EN', 'EN->ZH')
languageChosen.current(0)
'''
api = tk.StringVar()
apiChosen = ttk.Combobox(outmode_frame, width=8, textvariable=api)
apiChosen.grid(row = 3, column = 5, padx = 5, pady = 6, sticky = W)
apiChosen['values'] = ('接口选择', '有道翻译', '百度翻译', '腾讯翻译')
apiChosen.current(0)
'''
api = tk.StringVar()
api.set('有道翻译')
file = tk.StringVar()
fileChosen = ttk.Combobox(outmode_frame, width=8, textvariable=file)
fileChosen.grid(row = 3, column = 5, padx = 5, pady = 6, sticky = W)
fileChosen['values'] = ('输出选择', 'TXT', 'DOC')
fileChosen.current(0)
double_language = tk.IntVar()
ttk.Checkbutton(outmode_frame, text='双语输出', variable=double_language, onvalue=1, offvalue=0).grid(row = 4, column = 5, padx = 5, pady = 6, sticky = W)
ttk.Button(main_window, text = "提示", command = Tips).grid(row = 5, column = 0, pady=10)
ttk.Button(main_window, text = "关于", command = About).grid(row = 5, column = 1, pady=10)
ttk.Button(main_window, width = 25, text = "开始翻译", command = Translate).grid(row = 5, column = 2, columnspan = 2, padx = 15, pady=10)
out_frame = ttk.LabelFrame(main_window, text = '翻译文本输出')
out_frame.grid(row = 6, column = 0, columnspan = 5, padx=10, pady=10, sticky = W)
scr = scrolledtext.ScrolledText(out_frame, width = 56, height = 39, wrap=tk.WORD)
scr.grid(row = 6, column = 1, columnspan = 2, padx=5, pady=5)
pdf_frame = ttk.LabelFrame(main_window, text = 'PDF文件预览')
pdf_frame.grid(row = 0, column = 5, rowspan = 8, padx=10, pady=10, sticky = W)
cur_total = StringVar()
go_num = IntVar()
cur_total.set('0/0')
go_num.set(0)
ttk.Button(pdf_frame, text = "<", width = 2, command = Last_Page).grid(row = 0, column = 5, sticky = E)
ttk.Label(pdf_frame, textvariable = cur_total).grid(row = 0, column = 6)
ttk.Button(pdf_frame, text = ">", width = 2, command = Next_Page).grid(row = 0, column = 7, sticky = W)
go_entry = ttk.Entry(pdf_frame, width = 4, textvariable = go_num).grid(row = 0, column = 8, sticky = E)
ttk.Button(pdf_frame, text = "转到", width = 4, command = Go_Page).grid(row = 0, column = 9, sticky = W)
img = Image.open('.\\picture\\default.jpg')
img_temp = img.resize((440, 760), Image.ANTIALIAS)
image = ImageTk.PhotoImage(img_temp)
view = tk.Label(pdf_frame, width=440, height=780, image=image, compound='left')
view.grid(row=1, rowspan=8, column=5, columnspan=5)

main_window.mainloop()

